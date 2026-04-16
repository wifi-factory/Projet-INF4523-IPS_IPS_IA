from __future__ import annotations

import csv
import json
import shutil
from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
WEEK3_DIR = ROOT / "reports" / "week3_modelisation"
ASSETS_DIR = WEEK3_DIR / "report_assets"
SCREENSHOTS_SOURCE_DIR = Path(
    r"K:\4. UQO\04. INF4523 - Réseaux d'ordinateurs\07. IPS - IA orchestred - Project\04. IPS_IA_Semaine 3\screenshots_selection"
)
EXTERNAL_WEEK3_DIR = Path(
    r"K:\4. UQO\04. INF4523 - Réseaux d'ordinateurs\07. IPS - IA orchestred - Project\04. IPS_IA_Semaine 3"
)

CSV_PATH = WEEK3_DIR / "performance_modele_semaine3.csv"
SUMMARY_JSON_PATH = WEEK3_DIR / "semaine3_livrables_summary.json"
METADATA_PATH = ROOT / "models" / "random_forest_lab_v2_metadata.json"
SOURCE_INDEX_PATH = WEEK3_DIR / "code_source_module_detection.md"

DOCX_OUTPUT = WEEK3_DIR / "rapport_performance_semaine3_avec_screenshots.docx"
MARKDOWN_OUTPUT = WEEK3_DIR / "rapport_performance_semaine3_avec_screenshots.md"

SCREENSHOT_FILES = [
    "09. tshark_ping_normal.jpg",
    "10. csv_files.jpg",
    "13. script_build_datasetjpg.jpg",
    "14. Dataset_normal_screen.jpg",
    "15. Dataset_scan.jpg",
]

SCREENSHOT_CAPTIONS = {
    "09. tshark_ping_normal.jpg": (
        "Figure 1 - Capture reseau avec tshark sur trafic normal",
        "Cette capture montre une observation directe du trafic reseau brut. Elle rappelle que le pipeline du projet part bien d'une acquisition reelle avant transformation en variables exploitables par le modele.",
    ),
    "10. csv_files.jpg": (
        "Figure 2 - Fichiers CSV generes par le pipeline",
        "Cette figure illustre la production des fichiers intermediaires issus de la capture et du pretraitement. Elle permet de visualiser la transition entre trafic brut et donnees tabulaires.",
    ),
    "13. script_build_datasetjpg.jpg": (
        "Figure 3 - Script de construction du dataset",
        "Cette capture apporte une preuve visuelle de l'etape d'assemblage du dataset. Elle est utile pour montrer que la preparation des donnees repose sur un script explicite et reproductible.",
    ),
    "14. Dataset_normal_screen.jpg": (
        "Figure 4 - Extrait du dataset associe au trafic normal",
        "Cette figure montre un apercu de la structure tabulaire des donnees pour des observations normales. Elle aide a comprendre le format des entrees donnees au modele.",
    ),
    "15. Dataset_scan.jpg": (
        "Figure 5 - Extrait du dataset associe a un scenario de scan",
        "Cette figure montre un exemple de lignes liees a un comportement suspect de type scan. Elle permet d'illustrer que le dataset couvre aussi des scenarios offensifs.",
    ),
}


def ensure_dirs() -> None:
    WEEK3_DIR.mkdir(parents=True, exist_ok=True)
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)
    EXTERNAL_WEEK3_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, *, bold: bool = False):
    candidates = [
        r"C:\Windows\Fonts\calibri.ttf",
        r"C:\Windows\Fonts\arial.ttf",
    ]
    if bold:
        candidates = [
            r"C:\Windows\Fonts\calibrib.ttf",
            r"C:\Windows\Fonts\arialbd.ttf",
        ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def read_csv_rows(path: Path) -> list[dict[str, str]]:
    with path.open("r", encoding="utf-8", newline="") as handle:
        return list(csv.DictReader(handle))


def read_json(path: Path) -> dict[str, Any]:
    return json.loads(path.read_text(encoding="utf-8"))


def read_source_files() -> list[Path]:
    paths: list[Path] = []
    for line in SOURCE_INDEX_PATH.read_text(encoding="utf-8").splitlines():
        stripped = line.strip()
        if stripped.startswith("- `") and stripped.endswith("`"):
            paths.append(Path(stripped[3:-1]))
    return paths


def copy_screenshots() -> list[Path]:
    copied: list[Path] = []
    for filename in SCREENSHOT_FILES:
        source_path = SCREENSHOTS_SOURCE_DIR / filename
        target_path = ASSETS_DIR / filename
        shutil.copy2(source_path, target_path)
        copied.append(target_path)
    return copied


def to_float(value: str | None) -> float:
    if not value:
        return 0.0
    return float(value)


def find_row(rows: list[dict[str, str]], *, modele: str, split: str) -> dict[str, str]:
    for row in rows:
        if row["modele"] == modele and row["split"] == split:
            return row
    raise KeyError(f"Missing row for modele={modele}, split={split}")


def render_bar_chart(
    title: str,
    subtitle: str,
    categories: list[str],
    series: list[tuple[str, list[float], str]],
    output_path: Path,
    *,
    y_max: float | None = None,
) -> None:
    width = 1500
    height = 900
    left = 120
    right = width - 70
    top = 170
    bottom = height - 160
    plot_width = right - left
    plot_height = bottom - top
    bg = "#FFFFFF"
    text = "#202530"
    muted = "#5B6475"
    axis = "#C8D0DC"
    grid = "#E8EDF3"

    image = Image.new("RGB", (width, height), bg)
    draw = ImageDraw.Draw(image)
    title_font = load_font(34, bold=True)
    subtitle_font = load_font(18)
    label_font = load_font(18)
    value_font = load_font(16, bold=True)

    draw.text((left, 40), title, fill=text, font=title_font)
    draw.text((left, 90), subtitle, fill=muted, font=subtitle_font)

    if y_max is None:
        y_max = max((value for _, values, _ in series for value in values), default=1.0)
        y_max = max(y_max * 1.15, 1.0)

    for tick in range(6):
        tick_value = y_max * tick / 5
        y = bottom - int(plot_height * tick / 5)
        draw.line((left, y, right, y), fill=grid, width=1)
        draw.text((20, y - 10), f"{tick_value:.2f}", fill=muted, font=label_font)

    draw.line((left, top, left, bottom), fill=axis, width=2)
    draw.line((left, bottom, right, bottom), fill=axis, width=2)

    group_width = plot_width / max(len(categories), 1)
    inner_width = group_width * 0.72
    bar_width = inner_width / max(len(series), 1)

    for cat_idx, category in enumerate(categories):
        group_left = left + group_width * cat_idx + (group_width - inner_width) / 2
        center_x = left + group_width * cat_idx + group_width / 2
        draw.text((center_x - 35, bottom + 18), category, fill=text, font=label_font)
        for series_idx, (_, values, color) in enumerate(series):
            value = values[cat_idx]
            bar_left = group_left + series_idx * bar_width
            bar_right = bar_left + bar_width - 10
            bar_height = 0 if y_max == 0 else (value / y_max) * plot_height
            bar_top = bottom - bar_height
            draw.rounded_rectangle(
                (bar_left, bar_top, bar_right, bottom),
                radius=10,
                fill=color,
            )
            draw.text((bar_left + 3, bar_top - 24), f"{value:.3f}", fill=text, font=value_font)

    legend_y = height - 95
    legend_x = left
    for label, _, color in series:
        draw.rounded_rectangle((legend_x, legend_y, legend_x + 22, legend_y + 22), radius=5, fill=color)
        draw.text((legend_x + 34, legend_y - 1), label, fill=text, font=label_font)
        legend_x += 250

    image.save(output_path)


def add_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    for style_name, size in [("Heading 1", 16), ("Heading 2", 14), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = True


def add_title_page(document: Document) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("Projet INF4523\nRapport de performance - Semaine 3")
    run.bold = True
    run.font.size = Pt(20)

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run(
        "Modelisation IA, evaluation du modele et integration dans le pipeline IPS"
    )
    subtitle_run.font.size = Pt(12)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"Genere le {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    meta_run.italic = True
    document.add_section(WD_SECTION.NEW_PAGE)


def add_paragraph(document: Document, text: str) -> None:
    document.add_paragraph(text)


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        document.add_paragraph(item, style="List Bullet")


def add_table(document: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, header in enumerate(headers):
        hdr[idx].text = header
        hdr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def add_figure(document: Document, image_path: Path, caption_title: str, caption_text: str) -> None:
    document.add_picture(str(image_path), width=Inches(6.4))
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = caption.add_run(caption_title)
    title_run.bold = True
    document.add_paragraph(caption_text)


def build_markdown(
    *,
    metadata: dict[str, Any],
    summary: dict[str, Any],
    rows: list[dict[str, str]],
    screenshots: list[Path],
    chart_paths: list[Path],
) -> str:
    train = find_row(rows, modele="random_forest_lab_v2", split="train")
    validation = find_row(rows, modele="random_forest_lab_v2", split="validation")
    test = find_row(rows, modele="random_forest_lab_v2", split="test")
    baseline_validation = find_row(rows, modele="baseline_v1", split="validation")
    baseline_test = find_row(rows, modele="baseline_v1", split="test")

    lines: list[str] = []
    lines.append("# Rapport de performance Semaine 3 avec captures et graphes")
    lines.append("")
    lines.append("## 1. Introduction")
    lines.append(
        "Ce document presente une version enrichie du rapport de performance de la semaine 3. "
        "Il regroupe les metriques du modele retenu, des graphes de comparaison et plusieurs captures du pipeline de preparation des donnees."
    )
    lines.append("")
    lines.append("## 2. Modele retenu")
    lines.append(f"- Type de modele : {metadata['model_type']}")
    lines.append(f"- Classe positive : `{metadata['positive_label']}`")
    lines.append(f"- Vue de donnees : `{metadata['dataset_view_id']}`")
    lines.append(f"- Nombre de variables avant encodage : {len(metadata['input_columns_before_encoding'])}")
    lines.append("")
    lines.append("## 3. Tableau de synthese des performances")
    lines.append("")
    lines.append("| Split | Accuracy | Precision suspect | Recall suspect | F1 suspect | FPR | ROC-AUC |")
    lines.append("|---|---:|---:|---:|---:|---:|---:|")
    for row in (train, validation, test):
        lines.append(
            f"| {row['split']} | {float(row['accuracy']):.4f} | {float(row['precision_suspect']):.4f} | "
            f"{float(row['recall_suspect']):.4f} | {float(row['f1_suspect']):.4f} | "
            f"{float(row['false_positive_rate']):.4f} | {float(row['roc_auc_suspect']):.4f} |"
        )
    lines.append("")
    lines.append("## 4. Comparaison avec la baseline")
    lines.append(
        f"Sur validation, la baseline v1 atteint une accuracy de {float(baseline_validation['accuracy']):.4f}, "
        f"contre {float(validation['accuracy']):.4f} pour le modele retenu. "
        f"Sur test, le nouveau modele conserve aussi un net avantage, avec un FPR de {float(test['false_positive_rate']):.4f} "
        f"contre {float(baseline_test['false_positive_rate']):.4f} pour la baseline."
    )
    lines.append("")
    lines.append("## 5. Graphes de performance")
    for path in chart_paths:
        lines.append(f"- `{path.name}`")
    lines.append("")
    lines.append("## 6. Screenshots retenus")
    for path in screenshots:
        title, comment = SCREENSHOT_CAPTIONS[path.name]
        lines.append(f"### {title}")
        lines.append(comment)
        lines.append(f"Fichier : `{path.name}`")
        lines.append("")
    lines.append("## 7. Integration du module de detection")
    lines.append(
        "Le resume technique disponible confirme que le modele est integre dans le backend IPS : preparation des features, prediction, calcul de confiance, alerte et blocage."
    )
    lines.append("")
    lines.append("## 8. Remarques critiques")
    lines.append(
        "Les performances observees sont tres elevees. Elles doivent etre interpretees avec prudence dans la mesure ou les resultats proviennent d'un contexte de laboratoire controle."
    )
    lines.append(
        "Les artefacts visibles materialisent un pipeline RandomForest complet, mais pas d'implementation equivalente pour XGBoost, SVM ou Autoencoder."
    )
    lines.append("")
    lines.append("## 9. Conclusion")
    lines.append(
        "Le modele RandomForest retenu est non seulement entraine et evalue, mais aussi integre dans le pipeline de detection du backend. "
        "Le rapport enrichi produit ici rend cette semaine 3 plus presentable pour une remise universitaire."
    )
    lines.append("")
    lines.append(f"Note technique source : `{summary['deliverables']['performance_report_source']}`")
    lines.append("")
    return "\n".join(lines)


def main() -> int:
    ensure_dirs()
    rows = read_csv_rows(CSV_PATH)
    summary = read_json(SUMMARY_JSON_PATH)
    metadata = read_json(METADATA_PATH)
    copied_screenshots = copy_screenshots()
    source_files = read_source_files()

    train = find_row(rows, modele="random_forest_lab_v2", split="train")
    validation = find_row(rows, modele="random_forest_lab_v2", split="validation")
    test = find_row(rows, modele="random_forest_lab_v2", split="test")
    baseline_validation = find_row(rows, modele="baseline_v1", split="validation")
    baseline_test = find_row(rows, modele="baseline_v1", split="test")

    chart1 = ASSETS_DIR / "graph_metrics_par_split.png"
    render_bar_chart(
        "Comparaison des metriques du modele retenu par split",
        "Train, validation et test pour le RandomForest lab_v2",
        ["Accuracy", "Precision", "Recall", "F1"],
        [
            ("Train", [to_float(train["accuracy"]), to_float(train["precision_suspect"]), to_float(train["recall_suspect"]), to_float(train["f1_suspect"])], "#5B8DB8"),
            ("Validation", [to_float(validation["accuracy"]), to_float(validation["precision_suspect"]), to_float(validation["recall_suspect"]), to_float(validation["f1_suspect"])], "#8C6DD7"),
            ("Test", [to_float(test["accuracy"]), to_float(test["precision_suspect"]), to_float(test["recall_suspect"]), to_float(test["f1_suspect"])], "#4CA97B"),
        ],
        chart1,
        y_max=1.05,
    )

    chart2 = ASSETS_DIR / "graph_baseline_vs_modele.png"
    render_bar_chart(
        "Comparaison baseline v1 vs RandomForest lab_v2",
        "Performance sur validation et test",
        ["Val-Acc", "Val-F1", "Test-Acc", "Test-F1"],
        [
            ("Baseline v1", [to_float(baseline_validation["accuracy"]), to_float(baseline_validation["f1_suspect"]), to_float(baseline_test["accuracy"]), to_float(baseline_test["f1_suspect"])], "#D17C6A"),
            ("RandomForest lab_v2", [to_float(validation["accuracy"]), to_float(validation["f1_suspect"]), to_float(test["accuracy"]), to_float(test["f1_suspect"])], "#5B8DB8"),
        ],
        chart2,
        y_max=1.05,
    )

    chart3 = ASSETS_DIR / "graph_false_positive_rate.png"
    render_bar_chart(
        "Comparaison du taux de faux positifs",
        "Le FPR est un indicateur cle pour un systeme IPS",
        ["Val-B", "Val-RF", "Test-B", "Test-RF"],
        [
            ("FPR", [to_float(baseline_validation["false_positive_rate"]), to_float(validation["false_positive_rate"]), to_float(baseline_test["false_positive_rate"]), to_float(test["false_positive_rate"])], "#D65A6F"),
        ],
        chart3,
        y_max=max(to_float(baseline_validation["false_positive_rate"]), to_float(baseline_test["false_positive_rate"]), 0.65),
    )

    chart4 = ASSETS_DIR / "graph_roc_auc.png"
    render_bar_chart(
        "Comparaison du ROC-AUC",
        "Capacite de separation entre trafic normal et suspect",
        ["Val-B", "Val-RF", "Test-B", "Test-RF"],
        [
            ("ROC-AUC", [to_float(baseline_validation["roc_auc_suspect"]), to_float(validation["roc_auc_suspect"]), to_float(baseline_test["roc_auc_suspect"]), to_float(test["roc_auc_suspect"])], "#E0A84F"),
        ],
        chart4,
        y_max=1.05,
    )
    chart_paths = [chart1, chart2, chart3, chart4]

    markdown = build_markdown(
        metadata=metadata,
        summary=summary,
        rows=rows,
        screenshots=copied_screenshots,
        chart_paths=chart_paths,
    )
    MARKDOWN_OUTPUT.write_text(markdown, encoding="utf-8")

    document = Document()
    add_document_styles(document)
    add_title_page(document)

    document.add_heading("1. Introduction", level=1)
    add_paragraph(
        document,
        "Ce rapport presente une version enrichie de la semaine 3 du projet IPS base sur l'IA. "
        "Il rassemble les resultats de performance du modele retenu, plusieurs graphes de comparaison et des captures utiles pour documenter le pipeline de donnees.",
    )

    document.add_heading("2. Modele retenu et cadre d'entrainement", level=1)
    add_bullets(
        document,
        [
            f"Modele retenu : {metadata['model_type']}",
            f"Classe positive : {metadata['positive_label']}",
            f"Vue de donnees : {metadata['dataset_view_id']}",
            f"Nombre de variables avant encodage : {len(metadata['input_columns_before_encoding'])}",
            f"Train / validation / test : {metadata['training_summary']['train_rows']} / {metadata['training_summary']['validation_rows']} / {metadata['training_summary']['test_rows']} lignes",
        ],
    )
    add_paragraph(
        document,
        "Les artefacts visibles dans le depot montrent un pipeline RandomForest complet. "
        "Aucun pipeline equivalent pleinement materialise n'apparait pour XGBoost, SVM ou Autoencoder dans les fichiers examines.",
    )

    document.add_heading("3. Tableau de synthese des performances", level=1)
    add_table(
        document,
        ["Split", "Rows", "Accuracy", "Precision", "Recall", "F1", "FPR", "ROC-AUC"],
        [
            [
                train["split"],
                train["rows"],
                f"{to_float(train['accuracy']):.4f}",
                f"{to_float(train['precision_suspect']):.4f}",
                f"{to_float(train['recall_suspect']):.4f}",
                f"{to_float(train['f1_suspect']):.4f}",
                f"{to_float(train['false_positive_rate']):.4f}",
                f"{to_float(train['roc_auc_suspect']):.4f}",
            ],
            [
                validation["split"],
                validation["rows"],
                f"{to_float(validation['accuracy']):.4f}",
                f"{to_float(validation['precision_suspect']):.4f}",
                f"{to_float(validation['recall_suspect']):.4f}",
                f"{to_float(validation['f1_suspect']):.4f}",
                f"{to_float(validation['false_positive_rate']):.4f}",
                f"{to_float(validation['roc_auc_suspect']):.4f}",
            ],
            [
                test["split"],
                test["rows"],
                f"{to_float(test['accuracy']):.4f}",
                f"{to_float(test['precision_suspect']):.4f}",
                f"{to_float(test['recall_suspect']):.4f}",
                f"{to_float(test['f1_suspect']):.4f}",
                f"{to_float(test['false_positive_rate']):.4f}",
                f"{to_float(test['roc_auc_suspect']):.4f}",
            ],
        ],
    )
    add_paragraph(
        document,
        "Le modele conserve des scores tres eleves sur les trois splits. "
        "Le maintien de performances elevees sur validation et test suggere une bonne coherence avec les donnees disponibles dans le laboratoire.",
    )

    document.add_heading("4. Comparaison avec la baseline v1", level=1)
    add_table(
        document,
        ["Modele", "Split", "Accuracy", "Precision", "Recall", "F1", "FPR", "ROC-AUC"],
        [
            ["baseline_v1", "validation", f"{to_float(baseline_validation['accuracy']):.4f}", f"{to_float(baseline_validation['precision_suspect']):.4f}", f"{to_float(baseline_validation['recall_suspect']):.4f}", f"{to_float(baseline_validation['f1_suspect']):.4f}", f"{to_float(baseline_validation['false_positive_rate']):.4f}", f"{to_float(baseline_validation['roc_auc_suspect']):.4f}"],
            ["random_forest_lab_v2", "validation", f"{to_float(validation['accuracy']):.4f}", f"{to_float(validation['precision_suspect']):.4f}", f"{to_float(validation['recall_suspect']):.4f}", f"{to_float(validation['f1_suspect']):.4f}", f"{to_float(validation['false_positive_rate']):.4f}", f"{to_float(validation['roc_auc_suspect']):.4f}"],
            ["baseline_v1", "test", f"{to_float(baseline_test['accuracy']):.4f}", f"{to_float(baseline_test['precision_suspect']):.4f}", f"{to_float(baseline_test['recall_suspect']):.4f}", f"{to_float(baseline_test['f1_suspect']):.4f}", f"{to_float(baseline_test['false_positive_rate']):.4f}", f"{to_float(baseline_test['roc_auc_suspect']):.4f}"],
            ["random_forest_lab_v2", "test", f"{to_float(test['accuracy']):.4f}", f"{to_float(test['precision_suspect']):.4f}", f"{to_float(test['recall_suspect']):.4f}", f"{to_float(test['f1_suspect']):.4f}", f"{to_float(test['false_positive_rate']):.4f}", f"{to_float(test['roc_auc_suspect']):.4f}"],
        ],
    )
    add_paragraph(
        document,
        "La comparaison visible dans les artefacts montre un gain tres important par rapport a la baseline v1. "
        "Le point le plus marquant est la reduction du taux de faux positifs, qui etait eleve sur la baseline et devient tres faible sur le modele retenu.",
    )

    document.add_heading("5. Graphes de performance", level=1)
    graph_comments = {
        chart1.name: "Ce graphe compare Accuracy, Precision, Recall et F1-score sur train, validation et test. Il permet de verifier que les performances ne s'effondrent pas hors entrainement.",
        chart2.name: "Ce graphe montre la difference entre la baseline v1 et le RandomForest lab_v2 sur les splits validation et test. Il illustre l'amelioration academiquement la plus utile a commenter.",
        chart3.name: "Ce graphe met en avant le False Positive Rate. Dans un systeme IPS, cette mesure est centrale car un taux trop eleve degrade fortement l'utilite operationnelle du modele.",
        chart4.name: "Ce graphe montre le ROC-AUC sur validation et test. Il complete les autres metriques en illustrant la qualite globale de separation entre trafic normal et trafic suspect.",
    }
    for chart_path in chart_paths:
        add_figure(document, chart_path, chart_path.stem.replace("_", " "), graph_comments[chart_path.name])

    document.add_heading("6. Captures utiles du pipeline", level=1)
    for screenshot_path in copied_screenshots:
        title, comment = SCREENSHOT_CAPTIONS[screenshot_path.name]
        add_figure(document, screenshot_path, title, comment)

    document.add_heading("7. Integration dans le pipeline IPS", level=1)
    add_paragraph(
        document,
        "Le code source du module de detection confirme que le pipeline suit une logique claire : preparation des features, prediction du modele, calcul de confiance, declenchement d'alerte puis evaluation d'une decision de blocage.",
    )
    source_rows: list[list[str]] = []
    role_map = {
        "routes_detection.py": "Expose l'API de detection cote backend.",
        "detection_service.py": "Orchestre la detection d'un flux et la decision associee.",
        "model_service.py": "Charge le modele, verifie le contrat et produit les predictions.",
        "feature_service.py": "Prepare et convertit les features avant inference.",
        "schema_service.py": "Expose le contrat de schema et de metadonnees du modele.",
        "train_lab_v2_model.py": "Entraine le RandomForest et exporte le modele ainsi que ses rapports.",
    }
    for source_path in source_files:
        source_rows.append([source_path.name, role_map.get(source_path.name, "Role non detaille dans ce rapport."), str(source_path)])
    add_table(document, ["Fichier", "Role visible", "Chemin"], source_rows)

    document.add_heading("8. Discussion critique", level=1)
    add_paragraph(
        document,
        "Les scores observes sont tres eleves. Cela est positif pour le projet, mais ces chiffres doivent etre commentes avec prudence car ils proviennent d'un contexte de laboratoire controle. "
        "Dans un environnement reel plus heterogene, les performances pourraient etre moins stables."
    )
    add_paragraph(
        document,
        "Le rapport permet cependant de montrer un point important : le nouveau pipeline et le nouveau modele reduisent fortement les faux positifs par rapport a la baseline precedente, ce qui repond directement a l'un des problemes majeurs rencontres au debut du projet."
    )

    document.add_heading("9. Conclusion", level=1)
    add_paragraph(
        document,
        "La semaine 3 dispose maintenant d'un rapport de performance plus complet, combinant resultats numeriques, visualisations et captures du pipeline. "
        "Le RandomForest lab_v2 apparait comme un modele entraine, evalue et integre de facon coherente dans le backend IPS."
    )

    document.save(DOCX_OUTPUT)

    for output in [DOCX_OUTPUT, MARKDOWN_OUTPUT]:
        shutil.copy2(output, EXTERNAL_WEEK3_DIR / output.name)
    for chart_path in chart_paths:
        shutil.copy2(chart_path, EXTERNAL_WEEK3_DIR / chart_path.name)

    print(f"Generated: {DOCX_OUTPUT}")
    print(f"Generated: {MARKDOWN_OUTPUT}")
    for chart_path in chart_paths:
        print(f"Generated: {chart_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
