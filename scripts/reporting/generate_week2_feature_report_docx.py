from __future__ import annotations

import json
from collections import Counter
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[2]
WEEK2_DIR = ROOT / "reports" / "week2_preprocessing"
ASSETS_DIR = WEEK2_DIR / "report_assets"
DOCX_OUTPUT = WEEK2_DIR / "rapport_documentation_technique_features_semaine2.docx"
MARKDOWN_OUTPUT = WEEK2_DIR / "rapport_documentation_technique_features_semaine2.md"
METADATA_PATH = ROOT / "models" / "random_forest_lab_v2_metadata.json"
CATALOG_PATH = WEEK2_DIR / "feature_catalog.csv"
SUMMARY_PATH = WEEK2_DIR / "preprocessing_summary.json"
FLOW_AGGREGATION_PATH = ROOT / "backend" / "app" / "services" / "flow_aggregation_service.py"

PAGE_WIDTH_PX = 1500
PADDING_X = 42
PADDING_Y = 28
COLOR_BG = "#FFFFFF"
COLOR_TEXT = "#222222"
COLOR_MUTED = "#555555"
COLOR_BORDER = "#C7CCD3"
COLOR_HEADER = "#EEF3F8"
COLOR_GRID = "#E8EDF2"
COLOR_BAR = "#5B8DB8"
COLOR_BAR_ALT = "#8EB6D9"
COLOR_BAR_WARN = "#C27B4B"
COLOR_NORMAL = "#5B8DB8"
COLOR_SUSPECT = "#C65B5B"

CATEGORY_FR = {
    "network_identity": "Identification réseau",
    "temporal": "Temporelles",
    "volume": "Volume et taille",
    "packet_shape": "Statistiques",
    "tcp_flags": "TCP et protocole",
    "icmp": "TCP et protocole",
    "behavioral_window": "Comportementales",
    "payload_candidate": "Avancées",
}


def ensure_assets_dir() -> None:
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)


def load_font(size: int, *, bold: bool = False, mono: bool = False):
    candidates: list[str] = []
    if mono:
        candidates = [r"C:\Windows\Fonts\consola.ttf", r"C:\Windows\Fonts\cour.ttf"]
    elif bold:
        candidates = [
            r"C:\Windows\Fonts\calibrib.ttf",
            r"C:\Windows\Fonts\arialbd.ttf",
            r"C:\Windows\Fonts\timesbd.ttf",
        ]
    else:
        candidates = [
            r"C:\Windows\Fonts\calibri.ttf",
            r"C:\Windows\Fonts\arial.ttf",
            r"C:\Windows\Fonts\times.ttf",
        ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    if not text:
        return [""]
    words = text.split()
    if not words:
        return [text]
    lines = [words[0]]
    for word in words[1:]:
        trial = f"{lines[-1]} {word}"
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            lines[-1] = trial
        else:
            lines.append(word)
    return lines


def format_value(value: Any) -> str:
    if pd.isna(value):
        return ""
    if isinstance(value, float):
        if abs(value) >= 1000:
            return f"{value:,.2f}".replace(",", " ")
        return f"{value:.3f}".rstrip("0").rstrip(".")
    return str(value)


def load_data() -> tuple[dict[str, Any], pd.DataFrame, pd.DataFrame, pd.DataFrame, pd.DataFrame, dict[str, Any]]:
    metadata = json.loads(METADATA_PATH.read_text(encoding="utf-8"))
    train_df = pd.read_parquet(metadata["train_path"])
    validation_df = pd.read_parquet(metadata["validation_path"])
    test_df = pd.read_parquet(metadata["test_path"])
    catalog_df = pd.read_csv(CATALOG_PATH)
    summary = json.loads(SUMMARY_PATH.read_text(encoding="utf-8"))
    return metadata, train_df, validation_df, test_df, catalog_df, summary


def render_table_image(
    df: pd.DataFrame,
    title: str,
    output_path: Path,
    *,
    subtitle: str | None = None,
    max_rows: int | None = None,
) -> None:
    display_df = df.head(max_rows) if max_rows else df.copy()
    columns = list(display_df.columns)
    rows = [[format_value(value) for value in row] for row in display_df.to_numpy()]

    title_font = load_font(28, bold=True)
    subtitle_font = load_font(17)
    header_font = load_font(17, bold=True)
    body_font = load_font(15)

    probe = Image.new("RGB", (PAGE_WIDTH_PX, 200), COLOR_BG)
    draw = ImageDraw.Draw(probe)
    table_width = PAGE_WIDTH_PX - 2 * PADDING_X
    col_width = max(120, table_width // max(1, len(columns)))

    y = PADDING_Y
    y += draw.textbbox((0, 0), title, font=title_font)[3] + 10
    if subtitle:
        for line in wrap_text(draw, subtitle, subtitle_font, table_width):
            y += draw.textbbox((0, 0), line, font=subtitle_font)[3] + 4
        y += 8

    row_heights: list[int] = []
    for row in rows:
        max_lines = max(len(wrap_text(draw, cell, body_font, col_width - 16)) for cell in row)
        row_heights.append(max(32, max_lines * 22 + 10))
    height = y + 36 + sum(row_heights) + PADDING_Y

    image = Image.new("RGB", (PAGE_WIDTH_PX, height), COLOR_BG)
    draw = ImageDraw.Draw(image)
    y = PADDING_Y
    draw.text((PADDING_X, y), title, font=title_font, fill=COLOR_TEXT)
    y += draw.textbbox((0, 0), title, font=title_font)[3] + 10
    if subtitle:
        for line in wrap_text(draw, subtitle, subtitle_font, table_width):
            draw.text((PADDING_X, y), line, font=subtitle_font, fill=COLOR_MUTED)
            y += draw.textbbox((0, 0), line, font=subtitle_font)[3] + 4
        y += 8

    draw.rounded_rectangle(
        (PADDING_X, y, PADDING_X + table_width, y + 36),
        radius=10,
        fill=COLOR_HEADER,
        outline=COLOR_BORDER,
        width=1,
    )
    for idx, column in enumerate(columns):
        x = PADDING_X + idx * col_width
        draw.text((x + 8, y + 8), str(column), font=header_font, fill=COLOR_TEXT)
        if idx > 0:
            draw.line((x, y, x, y + 36 + sum(row_heights)), fill=COLOR_BORDER, width=1)
    y += 36

    for row_idx, row in enumerate(rows):
        row_height = row_heights[row_idx]
        fill_color = "#FFFFFF" if row_idx % 2 == 0 else "#F9FBFD"
        draw.rectangle(
            (PADDING_X, y, PADDING_X + table_width, y + row_height),
            fill=fill_color,
            outline=COLOR_BORDER,
            width=1,
        )
        for idx, cell in enumerate(row):
            x = PADDING_X + idx * col_width
            text_y = y + 6
            for line in wrap_text(draw, cell, body_font, col_width - 16):
                draw.text((x + 8, text_y), line, font=body_font, fill=COLOR_TEXT)
                text_y += 22
        y += row_height
    image.save(output_path)


def render_code_excerpt_image(
    source_path: Path,
    start_line: int,
    end_line: int,
    title: str,
    output_path: Path,
) -> None:
    excerpt = source_path.read_text(encoding="utf-8").splitlines()[start_line - 1 : end_line]
    code_lines = [f"{start_line + idx:>4}  {line.rstrip()}" for idx, line in enumerate(excerpt)]

    title_font = load_font(26, bold=True)
    body_font = load_font(16, mono=True)
    line_height = 24
    width = PAGE_WIDTH_PX
    height = PADDING_Y * 2 + 56 + len(code_lines) * line_height + 20

    image = Image.new("RGB", (width, height), "#FBFCFE")
    draw = ImageDraw.Draw(image)
    draw.text((PADDING_X, PADDING_Y), title, font=title_font, fill=COLOR_TEXT)
    panel_top = PADDING_Y + 44
    draw.rounded_rectangle(
        (PADDING_X, panel_top, width - PADDING_X, height - PADDING_Y),
        radius=14,
        fill="#0F1720",
        outline="#2E3A46",
        width=1,
    )
    y = panel_top + 16
    for line in code_lines:
        draw.text((PADDING_X + 18, y), line, font=body_font, fill="#D8E0EA")
        y += line_height
    image.save(output_path)


def render_bar_chart(
    data: list[tuple[str, float]],
    title: str,
    output_path: Path,
    *,
    subtitle: str | None = None,
    bar_color: str = COLOR_BAR,
) -> None:
    width = PAGE_WIDTH_PX
    height = 760
    chart_left = 150
    chart_right = width - 80
    chart_top = 150
    chart_bottom = height - 120
    gap = 18
    bar_width = max(24, int((chart_right - chart_left - gap * max(0, len(data) - 1)) / max(1, len(data))))

    title_font = load_font(28, bold=True)
    label_font = load_font(16)
    value_font = load_font(14, bold=True)

    image = Image.new("RGB", (width, height), COLOR_BG)
    draw = ImageDraw.Draw(image)
    draw.text((PADDING_X, PADDING_Y), title, font=title_font, fill=COLOR_TEXT)
    if subtitle:
        draw.text((PADDING_X, PADDING_Y + 38), subtitle, font=label_font, fill=COLOR_MUTED)

    max_value = max(value for _, value in data) if data else 1
    plot_height = chart_bottom - chart_top
    for tick in range(5):
        y = chart_bottom - int(plot_height * tick / 4)
        value = max_value * tick / 4
        draw.line((chart_left, y, chart_right, y), fill=COLOR_GRID, width=1)
        label = f"{value:,.0f}".replace(",", " ")
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.text((chart_left - bbox[2] - 12, y - 8), label, font=label_font, fill=COLOR_MUTED)

    for idx, (label, value) in enumerate(data):
        x = chart_left + idx * (bar_width + gap)
        bar_height = int((value / max_value) * plot_height) if max_value else 0
        y = chart_bottom - bar_height
        draw.rounded_rectangle((x, y, x + bar_width, chart_bottom), radius=10, fill=bar_color)
        vtxt = f"{value:,.0f}".replace(",", " ")
        vb = draw.textbbox((0, 0), vtxt, font=value_font)
        draw.text((x + (bar_width - vb[2]) / 2, y - 24), vtxt, font=value_font, fill=COLOR_TEXT)
        lb = draw.textbbox((0, 0), label, font=label_font)
        draw.text((x + max(0, (bar_width - lb[2]) / 2), chart_bottom + 10), label, font=label_font, fill=COLOR_TEXT)
    image.save(output_path)


def render_grouped_bar_chart(
    categories: list[str],
    series: list[tuple[str, list[float], str]],
    title: str,
    output_path: Path,
    *,
    subtitle: str | None = None,
) -> None:
    width = PAGE_WIDTH_PX
    height = 760
    chart_left = 130
    chart_right = width - 80
    chart_top = 150
    chart_bottom = height - 130

    title_font = load_font(28, bold=True)
    label_font = load_font(16)
    value_font = load_font(13, bold=True)

    image = Image.new("RGB", (width, height), COLOR_BG)
    draw = ImageDraw.Draw(image)
    draw.text((PADDING_X, PADDING_Y), title, font=title_font, fill=COLOR_TEXT)
    if subtitle:
        draw.text((PADDING_X, PADDING_Y + 38), subtitle, font=label_font, fill=COLOR_MUTED)

    max_value = max(max(values) for _, values, _ in series) if series else 1
    plot_height = chart_bottom - chart_top
    for tick in range(5):
        y = chart_bottom - int(plot_height * tick / 4)
        value = max_value * tick / 4
        draw.line((chart_left, y, chart_right, y), fill=COLOR_GRID, width=1)
        label = f"{value:,.0f}".replace(",", " ")
        bbox = draw.textbbox((0, 0), label, font=label_font)
        draw.text((chart_left - bbox[2] - 12, y - 8), label, font=label_font, fill=COLOR_MUTED)

    group_width = (chart_right - chart_left) / max(1, len(categories))
    bar_width = min(80, int((group_width - 30) / max(1, len(series))))

    for group_index, category in enumerate(categories):
        group_x = chart_left + group_index * group_width
        for series_index, (_, values, color) in enumerate(series):
            value = values[group_index]
            x = int(group_x + 15 + series_index * bar_width)
            bar_height = int((value / max_value) * plot_height) if max_value else 0
            y = chart_bottom - bar_height
            draw.rounded_rectangle((x, y, x + bar_width - 8, chart_bottom), radius=10, fill=color)
            draw.text((x, y - 22), f"{value:,.0f}".replace(",", " "), font=value_font, fill=COLOR_TEXT)
        cb = draw.textbbox((0, 0), category, font=label_font)
        draw.text((group_x + (group_width - cb[2]) / 2, chart_bottom + 12), category, font=label_font, fill=COLOR_TEXT)

    lx = chart_right - 240
    ly = chart_top - 70
    for label, _, color in series:
        draw.rounded_rectangle((lx, ly, lx + 16, ly + 16), radius=4, fill=color)
        draw.text((lx + 24, ly - 1), label, font=label_font, fill=COLOR_TEXT)
        ly += 28
    image.save(output_path)


def describe_feature(name: str, catalog_row: pd.Series | None = None) -> dict[str, str]:
    category = CATEGORY_FR.get(str(catalog_row["category"]) if catalog_row is not None else "", "Autres")
    if name == "protocol":
        return {
            "category": category,
            "description": "Indique le protocole du flux, par exemple TCP, UDP ou ICMP.",
            "method": "Visible dans le contrat final et encodé en One-Hot dans le pipeline scikit-learn.",
            "interest": "Permet de séparer des comportements de nature différente, comme HTTP, DNS ou ICMP.",
            "limits": "Le protocole seul ne suffit pas à conclure qu’un trafic est malveillant.",
        }
    if name in {"src_port", "dst_port"}:
        target = "source" if name == "src_port" else "destination"
        return {
            "category": category,
            "description": f"Représente le port {target} du flux.",
            "method": "La colonne est visible dans le dataset préparé et documentée comme issue du flow header.",
            "interest": "Aide à identifier le service impliqué ou le type d’échange observé.",
            "limits": "Un port sensible n’implique pas à lui seul un comportement suspect.",
        }
    if name == "duration_ms":
        return {
            "category": category,
            "description": "Mesure la durée du flux en millisecondes.",
            "method": "Le calcul est visible dans le service d’agrégation à partir du début et de la fin du flux.",
            "interest": "Aide à distinguer un échange très bref, un scan, une requête simple ou un transfert plus long.",
            "limits": "Une durée courte peut être normale ou suspecte selon le contexte.",
        }
    if name.startswith("packet_count_"):
        direction = "total" if name.endswith("total") else ("aller" if name.endswith("fwd") else "retour")
        return {
            "category": category,
            "description": f"Compte le nombre de paquets observés dans le sens {direction} du flux.",
            "method": "Le compteur est mis à jour directement lors de l’agrégation des paquets.",
            "interest": "Donne une mesure simple de l’intensité du trafic et de l’équilibre entre aller et retour.",
            "limits": "Cette variable ne décrit pas à elle seule la structure temporelle du flux.",
        }
    if name.startswith("byte_count_"):
        direction = "total" if name.endswith("total") else ("aller" if name.endswith("fwd") else "retour")
        return {
            "category": category,
            "description": f"Mesure le volume d’octets observé dans le sens {direction} du flux.",
            "method": "Le service d’agrégation additionne les tailles des paquets au fil du temps.",
            "interest": "Aide à distinguer des échanges très légers de flux plus volumineux.",
            "limits": "Le volume seul ne permet pas d’identifier la nature du trafic.",
        }
    if name.startswith("pkt_len_"):
        label = {
            "pkt_len_min": "taille minimale",
            "pkt_len_max": "taille maximale",
            "pkt_len_mean": "taille moyenne",
            "pkt_len_std": "dispersion des tailles",
        }[name]
        return {
            "category": category,
            "description": f"Décrit la {label} des paquets observés dans le flux.",
            "method": "La valeur est calculée pendant l’agrégation à partir des tailles successives.",
            "interest": "Ces mesures résument la forme du trafic et complètent les variables de volume.",
            "limits": "Elles sont moins informatives sur les flux très courts.",
        }
    if name.startswith("iat_"):
        label = {
            "iat_min_ms": "temps minimal entre deux paquets",
            "iat_max_ms": "temps maximal entre deux paquets",
            "iat_mean_ms": "temps moyen entre deux paquets",
            "iat_std_ms": "dispersion des temps entre paquets",
        }[name]
        return {
            "category": category,
            "description": f"Mesure le {label} dans un flux.",
            "method": "Le calcul est visible dans le code d’agrégation à partir des horodatages successifs.",
            "interest": "Ces variables décrivent le rythme du trafic et sa régularité.",
            "limits": "Leur lecture dépend du contexte et du nombre de paquets présents dans le flux.",
        }
    if name in {"syn_count", "ack_count", "rst_count", "fin_count", "psh_count"}:
        flag = name.replace("_count", "").upper()
        return {
            "category": category,
            "description": f"Compte le nombre d’occurrences du drapeau TCP {flag} dans le flux.",
            "method": "Le champ est prévu dans le service d’agrégation flow-level.",
            "interest": "En théorie, il peut aider à décrire la logique d’établissement ou de fermeture des connexions TCP.",
            "limits": "Dans les splits préparés visibles, cette variable existe mais reste nulle partout.",
        }
    if name in {"icmp_echo_req_count", "icmp_echo_reply_count"}:
        kind = "requêtes Echo Request" if name.endswith("req_count") else "réponses Echo Reply"
        return {
            "category": category,
            "description": f"Compte les {kind} observées dans le flux.",
            "method": "Le compteur est incrémenté selon le type ICMP visible dans le code d’agrégation.",
            "interest": "Permet de reconnaître et de quantifier les échanges de type ping.",
            "limits": "Le protocole ICMP reste très minoritaire dans le train observé.",
        }
    if name in {"connections_per_1s", "connections_per_5s"}:
        window = "une seconde" if name.endswith("1s") else "cinq secondes"
        return {
            "category": category,
            "description": f"Compte le nombre de connexions observées sur une fenêtre glissante de {window}.",
            "method": "Le calcul est visible dans `_add_window_features` à partir de `start_ts`, groupé par session et IP source.",
            "interest": "Utile pour repérer des rafales de connexions ou une activité anormalement dense.",
            "limits": "Cette variable dépend du rythme du lab et du périmètre observé.",
        }
    if name in {"distinct_dst_ports_per_5s", "distinct_dst_ips_per_5s"}:
        target = "ports de destination distincts" if "ports" in name else "adresses IP de destination distinctes"
        return {
            "category": category,
            "description": f"Compte les {target} sur une fenêtre de cinq secondes.",
            "method": "La valeur est calculée par ensemble dans la logique de fenêtre glissante du backend.",
            "interest": "Très utile pour représenter un comportement de scan multiports ou multicibles.",
            "limits": "Dans un petit lab, cette variable est limitée par le nombre de services et de cibles disponibles.",
        }
    if name == "icmp_packets_per_1s":
        return {
            "category": category,
            "description": "Compte le nombre de paquets ICMP observés sur une seconde.",
            "method": "Le calcul additionne les paquets des flux ICMP présents dans la fenêtre courte.",
            "interest": "Peut aider à représenter un burst ICMP.",
            "limits": "Cette variable n’est non nulle que sur très peu de lignes du train visible.",
        }
    if name == "failed_connection_ratio":
        return {
            "category": category,
            "description": "Ratio censé résumer les connexions TCP échouées.",
            "method": "Le calcul visible dans le backend repose sur les compteurs SYN, ACK et RST.",
            "interest": "En théorie, cette variable pourrait aider à repérer des connexions non abouties.",
            "limits": "Dans les jeux préparés visibles, cette colonne vaut 0 partout.",
        }
    if name == "payload_entropy":
        return {
            "category": "Avancées",
            "description": "Feature candidate qui viserait à mesurer la diversité des octets d’un payload.",
            "method": "Elle apparaît dans le catalogue comme candidate non retenue.",
            "interest": "Une telle variable pourrait être utile pour repérer certains trafics compressés, chiffrés ou artificiels.",
            "limits": "Cette information n’est pas calculée dans le contrat final déployé.",
        }
    return {
        "category": category,
        "description": "Cette information est visible comme colonne dans le dataset préparé.",
        "method": "Cette information n’est pas explicitement visible dans les éléments fournis.",
        "interest": "Cette information n’est pas explicitement visible dans les éléments fournis.",
        "limits": "Cette information n’est pas explicitement visible dans les éléments fournis.",
    }


def create_visual_assets(
    train_df: pd.DataFrame,
    validation_df: pd.DataFrame,
    test_df: pd.DataFrame,
    catalog_df: pd.DataFrame,
) -> dict[str, Path]:
    ensure_assets_dir()
    assets: dict[str, Path] = {}

    sample_cols = [
        "protocol",
        "src_port",
        "dst_port",
        "duration_ms",
        "packet_count_total",
        "byte_count_total",
        "connections_per_1s",
        "label_binary",
    ]
    assets["dataset_excerpt"] = ASSETS_DIR / "figure_dataset_excerpt.png"
    render_table_image(
        train_df[sample_cols].head(6),
        "Extrait réel du dataset préparé (train_balanced.parquet)",
        assets["dataset_excerpt"],
        subtitle="Les lignes ci-dessous proviennent directement du jeu d’entraînement utilisé par le modèle final.",
    )

    catalog_excerpt = pd.concat(
        [
            catalog_df.head(8),
            catalog_df[catalog_df["feature_name"].eq("payload_entropy")],
        ],
        ignore_index=True,
    )
    assets["feature_catalog_excerpt"] = ASSETS_DIR / "figure_feature_catalog_excerpt.png"
    render_table_image(
        catalog_excerpt[["feature_name", "category", "included_in_model", "source", "notes"]],
        "Extrait du catalogue réel des features",
        assets["feature_catalog_excerpt"],
        subtitle="Le catalogue synthétise l’origine des variables, leur famille logique et le cas particulier de payload_entropy.",
    )

    assets["code_excerpt"] = ASSETS_DIR / "figure_code_window_features.png"
    render_code_excerpt_image(
        FLOW_AGGREGATION_PATH,
        230,
        286,
        "Extrait du code d’agrégation montrant les fenêtres glissantes",
        assets["code_excerpt"],
    )

    family_counts = Counter(catalog_df["category"].tolist())
    assets["family_counts"] = ASSETS_DIR / "graph_feature_families.png"
    render_bar_chart(
        [(label, float(value)) for label, value in family_counts.items()],
        "Répartition des variables par famille",
        assets["family_counts"],
        subtitle="Comptage réalisé à partir du catalogue des features documenté dans le dépôt.",
    )

    assets["class_distribution"] = ASSETS_DIR / "graph_class_distribution.png"
    render_grouped_bar_chart(
        ["Train", "Validation", "Test"],
        [
            ("Normal", [float(train_df["label_binary"].eq("normal").sum()), float(validation_df["label_binary"].eq("normal").sum()), float(test_df["label_binary"].eq("normal").sum())], COLOR_NORMAL),
            ("Suspect", [float(train_df["label_binary"].eq("suspect").sum()), float(validation_df["label_binary"].eq("suspect").sum()), float(test_df["label_binary"].eq("suspect").sum())], COLOR_SUSPECT),
        ],
        "Répartition des classes dans les trois splits",
        assets["class_distribution"],
        subtitle="Le train est équilibré, tandis que validation et test contiennent un peu plus de flux suspects.",
    )

    assets["protocol_distribution"] = ASSETS_DIR / "graph_protocol_distribution.png"
    render_bar_chart(
        [(str(k), float(v)) for k, v in train_df["protocol"].value_counts().to_dict().items()],
        "Distribution du protocole dans le train",
        assets["protocol_distribution"],
        subtitle="Le protocole TCP domine nettement le train, suivi par UDP. ICMP reste marginal.",
        bar_color=COLOR_BAR_ALT,
    )

    assets["top_ports"] = ASSETS_DIR / "graph_top_destination_ports.png"
    render_bar_chart(
        [(str(int(k)), float(v)) for k, v in train_df["dst_port"].value_counts().head(10).to_dict().items()],
        "Top 10 des ports de destination dans le train",
        assets["top_ports"],
        subtitle="Ce graphe montre surtout les services les plus représentés dans le corpus utilisé pour l’entraînement.",
        bar_color=COLOR_BAR_WARN,
    )

    return assets


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = field_code
    fld_char_separate = OxmlElement("w:fldChar")
    fld_char_separate.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_separate)
    run._r.append(fld_char_end)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Title", 18), ("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.bold = True
        style.font.size = Pt(size)
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(6)

    if "Figure Caption" not in document.styles:
        style = document.styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Times New Roman"
        style.font.size = Pt(10)
        style.font.italic = True


def add_page_numbers(section) -> None:
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(paragraph, "PAGE")


def add_toc(document: Document) -> None:
    paragraph = document.add_paragraph()
    add_field(paragraph, 'TOC \\o "1-3" \\h \\z \\u')


def add_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_figure(document: Document, image_path: Path, caption: str, *, width_inches: float = 6.3) -> None:
    document.add_picture(str(image_path), width=Inches(width_inches))
    paragraph = document.add_paragraph(caption, style="Figure Caption")
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def set_cell_text(cell, text: str, *, bold: bool = False, size: int = 10) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP


def build_feature_table_rows(catalog_df: pd.DataFrame) -> list[list[str]]:
    rows = [[
        "Nom de la feature",
        "Catégorie",
        "Type",
        "Description",
        "Méthode d’obtention",
        "Utilité pour le modèle",
        "Remarques",
    ]]
    for _, row in catalog_df.iterrows():
        notes = describe_feature(str(row["feature_name"]), row)
        remarks = notes["limits"]
        if str(row["included_in_model"]).lower() == "false":
            remarks = f"{remarks} Feature candidate non retenue dans le contrat final."
        rows.append([
            str(row["feature_name"]),
            notes["category"],
            str(row["dtype_train"]),
            notes["description"],
            notes["method"],
            notes["interest"],
            remarks,
        ])
    return rows


def build_markdown(metadata: dict[str, Any], train_df: pd.DataFrame, validation_df: pd.DataFrame, test_df: pd.DataFrame, catalog_df: pd.DataFrame) -> str:
    lines: list[str] = []
    lines.append("# Documentation technique des features extraites")
    lines.append("")
    lines.append("## Résumé")
    lines.append("")
    lines.append("Ce rapport documente les variables extraites à partir du trafic réseau dans le cadre du prototype IPS basé sur l’IA. Il s’appuie uniquement sur les artefacts réellement visibles dans le dépôt : metadata du modèle, jeux préparés, catalogue des features et code d’agrégation flow-level.")
    lines.append("")
    lines.append(f"- Nombre de features retenues dans le contrat final : **{len(metadata['input_columns_before_encoding'])}**")
    lines.append(f"- Taille du train : **{len(train_df)}** lignes")
    lines.append(f"- Taille de la validation : **{len(validation_df)}** lignes")
    lines.append(f"- Taille du test : **{len(test_df)}** lignes")
    lines.append("")
    lines.append("## Limites visibles")
    lines.append("")
    lines.append("- `payload_entropy` apparaît comme feature candidate mais n’est pas implémentée dans le contrat final.")
    lines.append("- Les compteurs TCP (`syn_count`, `ack_count`, `rst_count`, `fin_count`, `psh_count`) et `failed_connection_ratio` sont présents dans le contrat, mais les statistiques visibles sur les splits préparés les montrent entièrement nuls.")
    lines.append("")
    lines.append("## Families")
    lines.append("")
    for category_key, group in catalog_df.groupby("category", sort=False):
        lines.append(f"- **{CATEGORY_FR.get(category_key, category_key)}** : {len(group)} variable(s)")
    lines.append("")
    return "\n".join(lines) + "\n"


def build_document(metadata: dict[str, Any], train_df: pd.DataFrame, validation_df: pd.DataFrame, test_df: pd.DataFrame, catalog_df: pd.DataFrame, summary: dict[str, Any], assets: dict[str, Path]) -> Document:
    document = Document()
    configure_styles(document)
    for section in document.sections:
        section.top_margin = Cm(2.2)
        section.bottom_margin = Cm(2.2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)
        add_page_numbers(section)

    title = document.add_paragraph()
    title.style = document.styles["Title"]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("Documentation technique des features extraites\n")
    title.add_run("Projet IPS basé sur l’IA – Semaine 2").font.size = Pt(14)
    intro = document.add_paragraph()
    intro.alignment = WD_ALIGN_PARAGRAPH.CENTER
    intro.add_run("Document généré à partir des artefacts réellement présents dans le dépôt.\n").italic = True
    intro.add_run("Version : 1er avril 2026")

    document.add_page_break()
    document.add_heading("Table des matières", level=1)
    add_toc(document)
    document.add_page_break()

    document.add_heading("1. Introduction", level=1)
    add_paragraph(document, "Ce rapport présente les features extraites à partir du trafic réseau dans le cadre du prototype d’IPS basé sur l’IA. L’objectif est de montrer comment des données brutes issues du laboratoire sont transformées en variables exploitables par un modèle de classification binaire.")
    add_paragraph(document, "Cette étape se situe entre la capture réseau et l’entraînement du modèle. Elle est essentielle, car un modèle d’IA ne peut pas apprendre correctement à partir de paquets bruts non structurés.")

    document.add_heading("2. Rôle du feature engineering dans le projet", level=1)
    add_paragraph(document, "Le feature engineering sert à convertir des échanges réseau en variables mesurables. Au lieu de considérer un paquet isolé, le projet travaille au niveau du flux. Cette approche permet de résumer la durée, le volume, le rythme et certains comportements visibles sur une courte fenêtre temporelle.")
    add_paragraph(document, "Cette transformation rend le trafic plus lisible pour le modèle. Elle aide aussi à représenter des comportements comme un scan, une rafale de connexions ou un échange applicatif plus classique.")

    document.add_heading("3. Source et nature des données", level=1)
    add_paragraph(document, "Les éléments visibles dans le dépôt montrent que les données proviennent de captures réseau du laboratoire converties ensuite en flux. Le script `scripts/lab_v2/pcap_to_flows.py` lit des fichiers PCAP via `tshark`, puis applique la logique d’agrégation du backend pour produire les lignes finales du dataset.")
    add_paragraph(document, "L’analyse est donc clairement réalisée au niveau flow-level. Chaque ligne des jeux `train_balanced.parquet`, `validation_clean.parquet` et `test_clean.parquet` décrit un flux agrégé, pas un paquet isolé.")
    add_paragraph(document, f"Les tailles observées sont les suivantes : train = {len(train_df)} lignes, validation = {len(validation_df)} lignes et test = {len(test_df)} lignes.")

    document.add_heading("4. Méthodologie d’extraction des features", level=1)
    add_paragraph(document, "La chaîne visible dans les scripts est la suivante : capture du trafic, lecture des paquets, agrégation bidirectionnelle en flux, ajout de variables de temps et de comportement, puis préparation finale pour le modèle.")
    add_paragraph(document, "Le prétraitement appliqué comprend la vérification du schéma attendu, la coercition stricte des types numériques, l’encodage de `protocol` en One-Hot et l’exclusion de colonnes de contexte comme `src_ip`, `dst_ip`, `capture_id`, `scenario_id` ou `severity`.")
    add_paragraph(document, "Dans les éléments fournis, la normalisation ne prend pas la forme d’une standardisation globale de type `StandardScaler` ou `MinMaxScaler`. Ici, la normalisation visible correspond surtout à un nettoyage des types et à un encodage cohérent des variables.")

    document.add_heading("5. Présentation détaillée des features extraites", level=1)
    add_paragraph(document, "Les features suivantes sont regroupées par familles logiques. Les descriptions restent volontairement simples, mais elles s’appuient sur des éléments réellement observables dans le code, les datasets et les fichiers de synthèse.")

    feature_order = list(metadata["input_columns_before_encoding"]) + ["payload_entropy"]
    groups: dict[str, list[str]] = {}
    for feature_name in feature_order:
        row = catalog_df.loc[catalog_df["feature_name"].eq(feature_name)]
        category_key = str(row.iloc[0]["category"]) if not row.empty else "other"
        groups.setdefault(CATEGORY_FR.get(category_key, category_key), []).append(feature_name)

    for index, (category_name, names) in enumerate(groups.items(), start=1):
        document.add_heading(f"5.{index} {category_name}", level=2)
        for feature_name in names:
            row = catalog_df.loc[catalog_df["feature_name"].eq(feature_name)]
            notes = describe_feature(feature_name, None if row.empty else row.iloc[0])
            paragraph = document.add_paragraph()
            paragraph.add_run(f"Nom de la feature : {feature_name}\n").bold = True
            paragraph.add_run(f"Catégorie : {notes['category']}\n")
            paragraph.add_run(f"Description simple : {notes['description']}\n")
            paragraph.add_run(f"Origine ou méthode de calcul visible : {notes['method']}\n")
            paragraph.add_run(f"Intérêt pour la détection : {notes['interest']}\n")
            paragraph.add_run(f"Limites ou remarques : {notes['limits']}")
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    document.add_heading("6. Tableau synthèse des features", level=1)
    add_paragraph(document, "Le tableau suivant reprend l’ensemble des features visibles dans le contrat final, ainsi que la feature candidate `payload_entropy` qui apparaît dans la documentation mais n’est pas retenue dans le modèle déployé.")
    table_rows = build_feature_table_rows(catalog_df)
    table = document.add_table(rows=1, cols=len(table_rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, value in enumerate(table_rows[0]):
        set_cell_text(table.rows[0].cells[idx], value, bold=True, size=9)
    for row in table_rows[1:]:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, size=8)

    document.add_page_break()
    document.add_heading("7. Captures d’écran commentées", level=1)
    add_paragraph(document, "Aucune capture d’écran interactive prise directement dans une VM du laboratoire n’a été ajoutée pendant cette génération. Les figures ci-dessous sont cependant des extraits visuels produits à partir d’artefacts réels du projet. Elles constituent une preuve visuelle utile sans inventer de contenu absent.")
    add_figure(document, assets["dataset_excerpt"], "Figure 1 – Extrait réel du dataset d’entraînement préparé. Cette figure montre des lignes issues de `train_balanced.parquet`, avec des variables de protocole, de ports, de durée, de volume et le label binaire.", width_inches=6.5)
    add_figure(document, assets["feature_catalog_excerpt"], "Figure 2 – Extrait du catalogue des features. On y voit la famille logique de plusieurs variables, leur origine visible et le cas particulier de `payload_entropy`, non retenue dans le contrat final.", width_inches=6.5)
    add_figure(document, assets["code_excerpt"], "Figure 3 – Extrait du code d’agrégation flow-level. Le bloc présenté montre la construction des variables sur fenêtres glissantes comme `connections_per_1s`, `connections_per_5s`, `distinct_dst_ports_per_5s` et `failed_connection_ratio`.", width_inches=6.5)

    document.add_heading("8. Graphes et interprétation visuelle", level=1)
    add_paragraph(document, "Les graphes ci-dessous sont générés à partir des données réellement disponibles dans les jeux préparés et dans le catalogue des features. Ils servent à mieux comprendre la structure du corpus et la place occupée par certaines familles de variables.")
    add_figure(document, assets["family_counts"], "Figure 4 – Répartition des variables par famille. Le contrat final est dominé par les variables de volume, temporelles et comportementales, ce qui confirme l’orientation flow-level du projet.")
    add_figure(document, assets["class_distribution"], "Figure 5 – Répartition des classes dans les trois splits. Le train est parfaitement équilibré, alors que validation et test contiennent légèrement plus de flux suspects.")
    add_figure(document, assets["protocol_distribution"], "Figure 6 – Distribution du protocole dans le train. Le trafic TCP domine nettement le corpus d’entraînement, suivi par UDP, tandis que l’ICMP reste marginal.")
    add_figure(document, assets["top_ports"], "Figure 7 – Top 10 des ports de destination dans le train. Les ports 80, 443, 53, 65000 et 22 ressortent clairement, ce qui reflète les scénarios web, DNS, scans et connexions administratives présents dans le laboratoire.")

    document.add_heading("9. Intérêt global des features pour le modèle IA", level=1)
    add_paragraph(document, "L’intérêt du jeu de features vient surtout de leur complémentarité. Certaines variables donnent le contexte réseau de base, d’autres mesurent la durée, le volume, la forme des paquets ou le rythme des connexions. Les variables comportementales, calculées sur fenêtres glissantes, sont particulièrement utiles pour représenter des scans, des bursts ou des répétitions rapides.")
    add_paragraph(document, "Cette combinaison est bien adaptée à une classification binaire normal/suspect, car elle ne repose pas sur un seul indice. Le modèle peut apprendre des motifs plus riches à partir de plusieurs colonnes prises ensemble.")

    document.add_heading("10. Limites et précautions", level=1)
    add_paragraph(document, "La première limite visible concerne `payload_entropy`. La variable apparaît bien dans la documentation technique, mais elle n’est pas calculée dans le contrat final déployé. Elle doit donc être présentée comme une piste envisagée, pas comme une feature réellement utilisée par le modèle final.")
    add_paragraph(document, "Une autre limite importante concerne les compteurs TCP `syn_count`, `ack_count`, `rst_count`, `fin_count`, `psh_count` ainsi que `failed_connection_ratio`. Ces colonnes existent dans le contrat, mais les statistiques descriptives visibles sur les trois splits les montrent entièrement nulles. Il faut donc éviter de leur attribuer un poids observé plus fort que ce que montrent réellement les données.")
    add_paragraph(document, "Enfin, certaines features dépendent fortement du contexte du laboratoire. Par exemple, `distinct_dst_ips_per_5s` reste mécaniquement faible dans un petit lab. Une variable prise seule ne suffit jamais à conclure qu’un trafic est malveillant.")

    document.add_heading("11. Conclusion", level=1)
    add_paragraph(document, "Le travail de feature engineering visible dans ce projet transforme correctement des captures réseau en un dataset exploitable par un modèle d’IA. La logique retenue reste simple, compréhensible et cohérente avec un prototype IPS flow-level.")
    add_paragraph(document, "Même si certaines pistes avancées comme l’entropie du payload n’ont pas été retenues, l’ensemble des variables extraites constitue une base solide pour distinguer le trafic normal du trafic suspect dans le laboratoire.")
    return document


def main() -> int:
    ensure_assets_dir()
    metadata, train_df, validation_df, test_df, catalog_df, summary = load_data()
    assets = create_visual_assets(train_df, validation_df, test_df, catalog_df)
    MARKDOWN_OUTPUT.write_text(build_markdown(metadata, train_df, validation_df, test_df, catalog_df), encoding="utf-8")
    document = build_document(metadata, train_df, validation_df, test_df, catalog_df, summary, assets)
    document.save(DOCX_OUTPUT)
    print(f"Generated markdown: {MARKDOWN_OUTPUT}")
    print(f"Generated docx: {DOCX_OUTPUT}")
    for name, path in assets.items():
        print(f"Generated asset [{name}]: {path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
