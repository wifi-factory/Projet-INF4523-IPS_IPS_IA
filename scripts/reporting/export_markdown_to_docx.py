from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt


def ensure_base_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    for style_name in ("Heading 1", "Heading 2", "Heading 3"):
        style = document.styles[style_name]
        style.font.name = "Calibri"

    document.styles["Heading 1"].font.size = Pt(16)
    document.styles["Heading 1"].font.bold = True
    document.styles["Heading 2"].font.size = Pt(14)
    document.styles["Heading 2"].font.bold = True
    document.styles["Heading 3"].font.size = Pt(12)
    document.styles["Heading 3"].font.bold = True

    if "Bullet List" not in [style.name for style in document.styles]:
        document.styles.add_style("Bullet List", WD_STYLE_TYPE.PARAGRAPH)


def append_markdown_line(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        return

    if stripped.startswith("### "):
        document.add_heading(stripped[4:].strip(), level=3)
        return
    if stripped.startswith("## "):
        document.add_heading(stripped[3:].strip(), level=2)
        return
    if stripped.startswith("# "):
        document.add_heading(stripped[2:].strip(), level=1)
        return
    if stripped.startswith("- "):
        document.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return

    document.add_paragraph(stripped)


def main() -> int:
    parser = argparse.ArgumentParser(description="Export a simple markdown-like file to DOCX.")
    parser.add_argument("--input", required=True, help="Input markdown file path.")
    parser.add_argument("--output", required=True, help="Output DOCX file path.")
    args = parser.parse_args()

    input_path = Path(args.input).resolve()
    output_path = Path(args.output).resolve()
    output_path.parent.mkdir(parents=True, exist_ok=True)

    content = input_path.read_text(encoding="utf-8")

    document = Document()
    ensure_base_styles(document)

    for line in content.splitlines():
        append_markdown_line(document, line)

    document.save(output_path)
    print(f"Generated DOCX: {output_path}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
