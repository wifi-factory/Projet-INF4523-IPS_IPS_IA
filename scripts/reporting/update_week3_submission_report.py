from __future__ import annotations

import os
import shutil
import tempfile
from pathlib import Path

from docx import Document


SOURCE_DOCX = Path(
    r"K:\4. UQO\04. INF4523 - Réseaux d'ordinateurs\07. IPS - IA orchestred - Project\04. IPS_IA_Semaine 3\Livrables\01. Rapport_Modelisation_IPS_Semaine3.docx"
)
TARGET_DOCX = Path(
    r"K:\4. UQO\04. INF4523 - Réseaux d'ordinateurs\07. IPS - IA orchestred - Project\04. IPS_IA_Semaine 3\Livrables\02. Rapport_Modelisation_IPS_Semaine3_corrige.docx"
)


REPLACEMENTS = {
    11: (
        "L’analyse ci-dessous s’appuie sur les artefacts techniques visibles pour la semaine 3 : le rapport de performance, le tableau CSV des métriques, le résumé JSON des livrables, les métadonnées du modèle et l’index du code source du module de détection. Lorsqu’un détail n’est pas explicitement observable dans ces éléments, il est signalé comme tel."
    ),
    18: (
        "L’intérêt de ce modèle dans un projet universitaire de détection d’intrusions tient à plusieurs points : il s’adapte bien aux données tabulaires, il reste relativement interprétable au niveau des grandes logiques de décision, et il permet souvent d’obtenir de bonnes performances sans pipeline trop complexe. Les métadonnées visibles permettent en outre d’identifier plusieurs hyperparamètres du modèle retenu : n_estimators=300, max_depth=20, min_samples_leaf=1 et max_features=sqrt. La stratégie exacte de pondération des classes n’est toutefois pas explicitement documentée dans les éléments fournis."
    ),
    23: (
        "La liste des 31 variables d’entrée est visible dans les métadonnées du modèle, mais la description détaillée de chaque feature, la méthode exacte d’équilibrage et le protocole complet de constitution du dataset ne sont pas entièrement documentés dans le rapport lui-même. Les captures du pipeline confirment toutefois que le travail part d’acquisitions réseau réelles, passe par des CSV intermédiaires, puis aboutit à un dataset tabulaire exploité par le modèle."
    ),
    26: (
        "Les hyperparamètres principaux du RandomForest sont visibles dans les métadonnées : n_estimators=300, max_depth=20, min_samples_leaf=1 et max_features=sqrt. En revanche, des paramètres comme min_samples_split ou class_weight ne sont pas explicitement renseignés dans les artefacts examinés."
    ),
    27: (
        "De la même manière, une éventuelle recherche d’hyperparamètres, une validation croisée systématique, un mécanisme de calibration des probabilités ou une stratégie avancée de sélection de variables ne sont pas explicitement visibles dans les éléments fournis. Sur ce point, le rapport doit rester mesuré : le modèle, le script d’entraînement et les paramètres principaux sont visibles, mais toute la procédure d’optimisation ne l’est pas."
    ),
    63: (
        "Les artefacts techniques de la semaine 3 montrent explicitement que le module de détection suit une logique de préparation des features, de prédiction, de calcul de confiance, de déclenchement d’alerte puis d’évaluation d’une décision de blocage. Cette chaîne est cohérente avec ce que l’on attend d’un backend IPS académique."
    ),
    68: (
        "Le code visible montre bien l’existence d’un calcul de confiance ainsi que de seuils d’alerte et de blocage, mais les valeurs opérationnelles précises de ces seuils ne sont pas entièrement documentées dans les éléments fournis. Il faut donc éviter de surinterpréter la mécanique interne. Ce que l’on peut affirmer, en revanche, c’est que la structure logicielle du backend prévoit bien un chemin complet entre l’entrée de features et une décision exploitable par le système."
    ),
    89: (
        "Plusieurs limites doivent donc être signalées honnêtement. D’abord, le contexte reste expérimental. Ensuite, la nature exacte de la baseline n’est pas explicitement documentée. Enfin, si les hyperparamètres principaux du modèle sont visibles, la procédure détaillée d’optimisation de l’entraînement et les valeurs opérationnelles précises des seuils backend ne sont pas entièrement documentées dans les éléments fournis."
    ),
}


def replace_paragraph_text(paragraph, new_text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(new_text)


def main() -> int:
    TARGET_DOCX.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp_source = Path(tmp_dir) / "source.docx"
        tmp_output = Path(tmp_dir) / "output.docx"
        shutil.copy2(SOURCE_DOCX, tmp_source)
        document = Document(tmp_source)

        for paragraph_number, new_text in REPLACEMENTS.items():
            paragraph = document.paragraphs[paragraph_number - 1]
            replace_paragraph_text(paragraph, new_text)

        document.save(tmp_output)
        shutil.copy2(tmp_output, TARGET_DOCX)

    print(f"Generated corrected report: {TARGET_DOCX}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
